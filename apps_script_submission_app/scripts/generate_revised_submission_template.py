from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "templates" / "innovation_submission_template_revised.docx"
PDF_PATH = BASE_DIR / "output" / "pdf" / "innovation_submission_template_revised.pdf"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(document):
    section = document.sections[0]
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    styles = document.styles
    styles["Normal"].font.name = "PingFang TC"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    styles["Normal"].font.size = Pt(10.5)


def add_docx_title(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第一屆聖五中學生創意設計挑戰賽\n作品介紹範本")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "PingFang TC"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")

    p2 = document.add_paragraph("此版本與提交網站及 AI 評改欄位一致。")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].italic = True
    p2.runs[0].font.size = Pt(10)


def add_basic_info_table(document):
    table = document.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Mm(35), Mm(135)]

    data = [
        ("班級", ""),
        ("組員", "學號1 姓名1；學號2 姓名2；學號3 姓名3"),
        ("作品名稱", ""),
    ]

    for row_idx, (label, value) in enumerate(data):
        for col_idx in range(2):
            cell = table.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.cell(row_idx, 0).text = label
        table.cell(row_idx, 1).text = value
        set_cell_shading(table.cell(row_idx, 0), "EAF2FF")


def add_section(document, title, prompt, hint=None, lines=3):
    p = document.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)

    prompt_p = document.add_paragraph(prompt)
    prompt_p.paragraph_format.space_after = Pt(2)
    if hint:
        hint_p = document.add_paragraph(hint)
        hint_p.runs[0].italic = True
        hint_p.runs[0].font.size = Pt(9.5)
        hint_p.paragraph_format.space_after = Pt(3)

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = "\n" * lines


def build_docx():
    document = Document()
    set_doc_defaults(document)
    add_docx_title(document)
    add_basic_info_table(document)

    sections = [
        (
            "一、問題說明",
            "請在 200 字內說明你們想解決的真實問題。",
            "建議交代問題出現在哪裡、影響誰，以及為何值得處理。",
            4,
        ),
        (
            "二、設計理念",
            "請在 500 字內說明作品的核心想法、設計方向與價值。",
            "可寫為何想到這個方案，與現有做法相比有何不同。",
            5,
        ),
        (
            "三、解決方案細節",
            "請在 1000 字內說明作品如何運作、如何使用，以及是否可行。",
            "可包括流程、功能、材料、技術、成本或實施步驟。",
            6,
        ),
        (
            "四、人工智能使用情況",
            "請如實說明本次設計有否使用 AI，以及使用的平台與用途。",
            "如未有使用，請直接寫「本作品未有使用 AI 工具」。",
            4,
        ),
        (
            "五、科學海報",
            "另交 1 份 JPG 海報，總結作品名稱、問題、設計理念、解決方案與價值。",
            "海報用作視覺總結，不代替文字欄位。",
            3,
        ),
    ]

    for title, prompt, hint, lines in sections:
        add_section(document, title, prompt, hint, lines)

    document.add_paragraph("")
    note = document.add_paragraph("註：本範本已對齊最新參考資料與提交系統。")
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(9.5)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(DOCX_PATH))


def build_pdf():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CJKTitle",
            fontName="STSong-Light",
            fontSize=18,
            leading=24,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#16324f"),
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKSub",
            fontName="STSong-Light",
            fontSize=10,
            leading=14,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#5b6573"),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKHead",
            fontName="STSong-Light",
            fontSize=12,
            leading=16,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#16324f"),
            spaceBefore=6,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKBody",
            fontName="STSong-Light",
            fontSize=10.2,
            leading=15,
            alignment=TA_LEFT,
            textColor=colors.black,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CJKHint",
            fontName="STSong-Light",
            fontSize=9.2,
            leading=13,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#5b6573"),
            spaceAfter=3,
        )
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )

    story = [
        Paragraph("第一屆聖五中學生創意設計挑戰賽", styles["CJKTitle"]),
        Paragraph("作品介紹範本", styles["CJKTitle"]),
        Paragraph("此版本與提交網站及 AI 評改欄位一致。", styles["CJKSub"]),
    ]

    basic = Table(
        [
            ["班級", ""],
            ["組員", "學號1 姓名1；學號2 姓名2；學號3 姓名3"],
            ["作品名稱", ""],
        ],
        colWidths=[42 * mm, 130 * mm],
    )
    basic.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("LEADING", (0, 0), (-1, -1), 14),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EAF2FF")),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#B9C7DB")),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([basic, Spacer(1, 8)])

    sections = [
        ("一、問題說明", "請在 200 字內說明你們想解決的真實問題。", "建議交代問題位置、影響對象與重要性。"),
        ("二、設計理念", "請在 500 字內說明作品的核心想法、設計方向與價值。", "可寫為何想到這個方案，以及與現有做法的差異。"),
        ("三、解決方案細節", "請在 1000 字內說明作品如何運作、如何使用，以及是否可行。", "可包括流程、功能、材料、技術、成本或實施步驟。"),
        ("四、人工智能使用情況", "請如實說明本次設計有否使用 AI，以及使用的平台與用途。", "如未有使用，請直接寫「本作品未有使用 AI 工具」。"),
        ("五、科學海報", "另交 1 份 JPG 海報，總結作品名稱、問題、設計理念、解決方案與價值。", "海報用作視覺總結，不代替文字欄位。"),
    ]

    for title, prompt, hint in sections:
        story.append(Paragraph(title, styles["CJKHead"]))
        story.append(Paragraph(prompt, styles["CJKBody"]))
        story.append(Paragraph(hint, styles["CJKHint"]))
        filler = Table([["\n\n\n"]], colWidths=[172 * mm])
        filler.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                    ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#B9C7DB")),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.extend([filler, Spacer(1, 6)])

    PDF_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
